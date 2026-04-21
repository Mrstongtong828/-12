"""
BLOB/图片/文档扫描器 v5.1 (架构师优化版)。

优化重点(联动 patterns.py & encoded.py):
  1. OCR 多行文本拼接: 解决证件照姓名/地址被切行。
  2. 身份证号 X/N 容错: 解决 OCR 对末位 X 的空格截断及形近字误识。
  3. 姓名纠错映射: 针对小字体 OCR 常犯的字形错误进行逆向修复。
  4. 多路径并行扫描: 原文/数字修复/汉字修复/全修复，四路合并去重，极限拉升 Recall。
  5. 兼容 encoded.py 的入口调用逻辑。
"""
import re
import io

from core.patterns import extract_sensitive_from_value
from core.config import SENSITIVE_LEVEL_MAP
from scanners.ocr_client import get_ocr_text, ocr_disabled

# ── 文档解析参数 ─────────────────────────────────────────────────
PDF_MAX_PAGES = 20
PDF_OCR_DPI = 150
DOC_TEXT_MAX_LEN = 500_000

# ── 延迟加载容器 ─────────────────────────────────────────────────
_PYMUPDF_TRIED = False
_PYMUPDF = None
_DOCX_TRIED = False
_DOCX = None
_OPENPYXL_TRIED = False
_OPENPYXL = None


def _get_pymupdf():
    global _PYMUPDF_TRIED, _PYMUPDF
    if not _PYMUPDF_TRIED:
        _PYMUPDF_TRIED = True
        try:
            import fitz
            _PYMUPDF = fitz
            print("[INFO] PyMuPDF 可用, PDF 将直接抽文本")
        except Exception as e:
            print(f"[INFO] PyMuPDF 不可用, PDF 路径将退化或返空: {e}")
    return _PYMUPDF


def _get_docx():
    global _DOCX_TRIED, _DOCX
    if not _DOCX_TRIED:
        _DOCX_TRIED = True
        try:
            import docx
            _DOCX = docx
            print("[INFO] python-docx 可用")
        except Exception:
            pass
    return _DOCX


def _get_openpyxl():
    global _OPENPYXL_TRIED, _OPENPYXL
    if not _OPENPYXL_TRIED:
        _OPENPYXL_TRIED = True
        try:
            import openpyxl
            _OPENPYXL = openpyxl
            print("[INFO] openpyxl 可用")
        except Exception:
            pass
    return _OPENPYXL


# ═══════════════════════════════════════════════════════════════
# 数字类 OCR 纠错 (字母→数字)
# ═══════════════════════════════════════════════════════════════
OCR_CORRECTION_MAP = {
    'O': '0', 'o': '0', 'Q': '0', 'D': '0',
    'I': '1', 'l': '1', 'i': '1', '|': '1', 'L': '1',
    'Z': '2', 'z': '2',
    'S': '5', 's': '5',
    'G': '6', 'b': '6',
    'T': '7',
    'B': '8',
    'g': '9', 'q': '9',
}
# 匹配至少6个字符的数字疑似串 (加入 L 大写以增强银行卡/身份证召回)
_DIGIT_LIKE = re.compile(r"[0-9OoQDIlLiZzSsGbTBgq|]{6,}")

# [新增] 纯数字长窗口正则:用于 OCR 后仍有粘连的长数字块兜底(银行卡/身份证)
_LONG_DIGIT_WINDOW = re.compile(r"\d{13,20}")

def correct_ocr_text(text: str) -> str:
    """
    数字区段 OCR 纠错,特殊处理:
      - 18位 + 末位 X/x → 末位保留为 'X',前 17 位纠错
      - 18位 + 末位 N/n → 猜测为身份证 X 末位误识,改成 'X'
      - 其他长度段全部按 map 纠错
    """
    def _correct_segment(m):
        seg = m.group()
        # 18位长度: 处理身份证 X/N 末位误识
        if len(seg) == 18:
            last = seg[-1]
            if last in "XxNn":
                front = "".join(OCR_CORRECTION_MAP.get(c, c) for c in seg[:17])
                return front + 'X'
        # 通用数字纠错
        return "".join(OCR_CORRECTION_MAP.get(c, c) for c in seg)

    return _DIGIT_LIKE.sub(_correct_segment, text)


# ═══════════════════════════════════════════════════════════════
# 数字-数字 混淆枚举器 (校验位反馈驱动)
# ═══════════════════════════════════════════════════════════════
# Why: OCR_CORRECTION_MAP 只能修"字母→数字"方向的错误。PaddleOCR 常见的
#      "真 6 识成 0"、"真 8 识成 3" 等纯数字混淆已绕过字母阶段,
#      正则抓到的 digits_only 看起来是合法数字串但校验位过不去 —
#      此时按字形相近图谱做有界枚举,交给 GB11643 / Luhn / 模板正则反验。
# 图谱来源:常见 OCR 混淆矩阵 + 字形分析。每对最多 3 个邻居,控制组合爆炸。
DIGIT_CONFUSION = {
    '0': ('6', '8', '9'),
    '1': ('7',),
    '2': (),
    '3': ('8', '5'),
    '4': ('9',),
    '5': ('6', '3', '8'),
    '6': ('0', '5', '8'),
    '7': ('1',),
    '8': ('0', '3', '6'),
    '9': ('0', '4'),
}
# 高混淆数字:这几位最常被 OCR 读错,BFS 时优先从这些位置开始枚举
_HIGH_SUSP_DIGITS = frozenset('0135689')


def _enumerate_digit_fixes(digits: str, validator, max_edits: int = 2,
                            max_candidates: int = 128):
    """
    BFS 枚举数字串的 1~max_edits 次 digit→digit 混淆替换,返回首个通过
    validator 的版本;找不到返回 None。

    设计:
      - 仅在首轮 validator(digits) 失败时调用(外部保证)。
      - 位置按"嫌疑分"排序,高混淆位优先枚举,通常 1-edit 内命中。
      - max_candidates 硬顶 BFS 总探索量,防止组合爆炸。
      - tried 集合去重,避免 BFS 不同路径产出同一候选重复验证。
    """
    if not digits or not digits.isdigit():
        return None
    if validator(digits):
        return digits

    positions = sorted(range(len(digits)),
                        key=lambda i: 0 if digits[i] in _HIGH_SUSP_DIGITS else 1)
    tried = {digits}
    frontier = [(digits, 0)]
    head = 0  # 用索引代替 pop(0),避免 list.pop(0) 的 O(n) 开销

    while head < len(frontier) and len(tried) < max_candidates:
        s, depth = frontier[head]
        head += 1
        if depth >= max_edits:
            continue
        for pos in positions:
            for cand in DIGIT_CONFUSION.get(s[pos], ()):
                new_s = s[:pos] + cand + s[pos + 1:]
                if new_s in tried:
                    continue
                tried.add(new_s)
                if validator(new_s):
                    return new_s
                if len(tried) >= max_candidates:
                    return None
                frontier.append((new_s, depth + 1))
    return None


# ═══════════════════════════════════════════════════════════════
# 姓名汉字 OCR 纠错 (OCR 字形相近误识)
# ═══════════════════════════════════════════════════════════════
NAME_CHAR_FIX = {
    '郡': '邵',
    '偌': '储',
    '窃': '窦',
    '沟': '汲',
    '货': '贡',
    '历': '厉',
    '方': '万',  # '方'本身也是姓, 但在特定极高误报的业务字典中退化处理
    '姊': '姚',
    '妹': '姊',
    '漷': '漕',
    '潰': '漕',
    '湖': '郜',
    '郭': '郜',  # 郜 与 郭 在小字体下极像
}

def _name_char_retry_text(text: str) -> str:
    """
    姓名字形纠错: 仅对长度 2-4 的纯汉字 token 做首字映射。
    不改变原文, 返回纠错后的副本供二次扫描。
    """
    if not text:
        return text

    def _fix_token(m):
        tok = m.group()
        if len(tok) < 2 or len(tok) > 4:
            return tok
        first = tok[0]
        fixed = NAME_CHAR_FIX.get(first)
        if fixed and fixed != first:
            return fixed + tok[1:]
        return tok

    return re.sub(r"[\u4e00-\u9fa5]{2,4}", _fix_token, text)


# ═══════════════════════════════════════════════════════════════
# OCR 文本预处理 (清洗降噪)
# ═══════════════════════════════════════════════════════════════
_OCR_SPACE_IN_LONGNUM = re.compile(r"(\d)\s+(\d)")
_OCR_LABEL_COLON = re.compile(
    r"(姓名|名字|客户名|联系人|身份证号|身份证|手机|电话|联系电话|"
    r"银行卡号|卡号|病历号|就诊号|住址|地址|邮箱|邮件|持卡人|签发|"
    r"发卡行|开户行|户籍|籍贯|民族|出生|身份)\s*[:\uff1a]\s*",
)

_OCR_CERT_PREFIX_SPACE = re.compile(r"\b(MR|YB|mR|mr|Mr)\s*(\d)")
_OCR_MR_NORMALIZE = re.compile(r"\b[mM][rR](?=\d)")
_OCR_YB_NORMALIZE = re.compile(r"\b[yY][bB](?=\d)")
_OCR_CJK_SPACE = re.compile(r"([\u4e00-\u9fa5])[ \t]{1,2}([\u4e00-\u9fa5])")
_OCR_ID_X_SPACE = re.compile(r"(\d{17})\s+([XxNn])(?![\w\d])")

# 地址跨行吸附：保留省市区锚点
_ADDR_ADMIN_TAIL = re.compile(
    r"([\u4e00-\u9fa5]*(?:省|市|区|县|镇|乡|街道|路|街|弄|巷))\n"
    r"([\u4e00-\u9fa5\d][\u4e00-\u9fa5\d]*)"
)

def _preprocess_ocr_text(text: str) -> str:
    """
    OCR 文本预处理流水线：合并碎片、修复空格、统一病历前缀
    """
    if not text:
        return text

    # 1) Label 规整
    text = _OCR_LABEL_COLON.sub(lambda m: m.group(1) + ":", text)

    # 2) MR/YB 大小写归一
    text = _OCR_MR_NORMALIZE.sub("MR", text)
    text = _OCR_YB_NORMALIZE.sub("YB", text)

    # 3) 数字空格合并 (应对如 6222 0214 格式的银行卡)
    for _ in range(3):
        new = _OCR_SPACE_IN_LONGNUM.sub(r"\1\2", text)
        if new == text: break
        text = new

    # 4) 证件号前缀空格 (MR 123456 -> MR123456)
    text = _OCR_CERT_PREFIX_SPACE.sub(r"\1\2", text)

    # 5) 身份证末位空格截断修复
    text = _OCR_ID_X_SPACE.sub(r"\1\2", text)

    # 6) 汉字间单空格合并 (如 "张 三")
    for _ in range(2):
        new = _OCR_CJK_SPACE.sub(r"\1\2", text)
        if new == text: break
        text = new

    # 7) 地址跨行合并 (不丢失核心锚点)
    text_joined = text
    for _ in range(3):
        new = _ADDR_ADMIN_TAIL.sub(r"\1\2", text_joined)
        if new == text_joined: break
        text_joined = new

    # 8) 多行短汉字合并 (专门应对竖排版姓名被切成两行)
    lines = text_joined.split("\n")
    merged_lines = []
    i = 0
    while i < len(lines):
        cur = lines[i].strip()
        if (i + 1 < len(lines)
                and 0 < len(cur) <= 5
                and all('\u4e00' <= c <= '\u9fa5' for c in cur)):
            nxt = lines[i + 1].strip()
            if (0 < len(nxt) <= 5
                    and all('\u4e00' <= c <= '\u9fa5' for c in nxt)
                    and len(cur) + len(nxt) <= 6):
                merged_lines.append(cur + nxt)
                i += 2
                continue
        merged_lines.append(lines[i])
        i += 1
    
    return "\n".join(merged_lines)


# ═══════════════════════════════════════════════════════════════
# finding 构造与文件嗅探
# ═══════════════════════════════════════════════════════════════
def _make_finding(db_type, db_name, table_name, field_col_name,
                  record_id, sensitive_type, extracted_value):
    level = SENSITIVE_LEVEL_MAP.get(sensitive_type, "L3")
    return {
        "db_type": db_type,
        "db_name": db_name,
        "table_name": table_name,
        "field_name": field_col_name,
        "record_id": record_id,
        "data_form": "binary_blob",
        "sensitive_type": sensitive_type,
        "sensitive_level": level,
        "extracted_value": extracted_value,
    }

def _normalize_to_bytes(blob):
    if blob is None: return None
    if isinstance(blob, memoryview): return bytes(blob)
    if isinstance(blob, bytearray): return bytes(blob)
    if isinstance(blob, bytes): return blob
    if isinstance(blob, str):
        s = blob.strip()
        # ① PostgreSQL \x 十六进制
        if s.startswith(r'\x') and len(s) > 4 and (len(s) - 2) % 2 == 0:
            try: return bytes.fromhex(s[2:])
            except ValueError: return None
        # ② 裸十六进制(无 \x 前缀),长度 > 100 且全 hex
        if len(s) > 100 and len(s) % 2 == 0 and all(c in "0123456789abcdefABCDEF" for c in s):
            try: return bytes.fromhex(s)
            except ValueError: pass
    return None


def _unwrap_encoded_blob(data: bytes) -> bytes:
    """
    BLOB 套娃解码兜底:部分题目 BLOB 实际存的不是原始图片,而是
      - base64 → 图片
      - hex → base64 → 图片
    等嵌套形式。通过尝试 base64/hex 单层解码,若解码结果以图片 magic 起头,
    则返回解码后的字节,让 _sniff_file_type 识别为 image。
    仅做一层解码,避免过度尝试拖慢 OCR 主流程。
    """
    if not data or len(data) < 80:
        return data

    # 只处理"ASCII-only"的 BLOB(原始图片字节含大量非 ASCII,绝不会误触)
    try:
        text = data.decode("ascii", errors="strict").strip()
    except UnicodeDecodeError:
        return data

    # 尝试 base64
    if len(text) > 80 and all(c in
            "ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz0123456789+/=-_\r\n"
            for c in text):
        try:
            import base64 as _b64
            compact = text.replace("\n", "").replace("\r", "")
            padded = compact + "=" * (-len(compact) % 4)
            if "-" in compact or "_" in compact:
                raw = _b64.urlsafe_b64decode(padded)
            else:
                raw = _b64.b64decode(padded, validate=False)
            if any(raw.startswith(m) for m in (
                    b'\xff\xd8', b'\x89PNG', b'GIF8', b'BM', b'RIFF', b'%PDF')):
                return raw
        except Exception:
            pass

    # 尝试 hex
    if len(text) % 2 == 0 and all(c in "0123456789abcdefABCDEF" for c in text):
        try:
            raw = bytes.fromhex(text)
            if any(raw.startswith(m) for m in (
                    b'\xff\xd8', b'\x89PNG', b'GIF8', b'BM', b'RIFF', b'%PDF')):
                return raw
        except Exception:
            pass

    return data


def _sniff_file_type(data: bytes) -> str:
    if not data or len(data) < 8: return "unknown"
    if data[:4] == b'%PDF': return "pdf"
    if data[:2] == b'\xff\xd8' or data[:4] == b'\x89PNG' or data[:4] == b'GIF8' or data[:2] == b'BM' or data[:4] == b'RIFF':
        return "image"
    if data[:4] == b'PK\x03\x04':
        head = data[:4096]
        if b'word/document.xml' in head or b'word/' in head or b'[Content_Types].xml' in head:
            return "docx"
        if b'xl/workbook.xml' in head or b'xl/' in head:
            return "xlsx"
    return "unknown"


def _bytes_as_text_fallback(data: bytes) -> str:
    """
    BLOB 无法识别文件类型 / OCR 返回空时的最后兜底:
    将原始字节按多种编码尝试解码,提取出其中的 ASCII/UTF-8 可打印文本,
    交给正则管线做一次扫描。覆盖"BLOB 里其实塞了明文 + 少量控制字符"场景。
    """
    if not data or len(data) < 16:
        return ""
    texts = []
    for enc in ("utf-8", "gbk", "latin-1"):
        try:
            t = data.decode(enc, errors="ignore")
            # 过滤不可打印字符,只保留常规文本 + 中文
            clean = "".join(
                c for c in t
                if c.isprintable() or c in "\n\r\t"
                or '\u4e00' <= c <= '\u9fa5'
            )
            if len(clean) > 20:
                texts.append(clean)
                break
        except Exception:
            continue
    return "\n".join(texts)[:DOC_TEXT_MAX_LEN]


# ═══════════════════════════════════════════════════════════════
# PDF / DOCX / XLSX 文本抽取
# ═══════════════════════════════════════════════════════════════
def _extract_pdf_text(data: bytes) -> str:
    fitz = _get_pymupdf()
    if fitz is None: return ""
    try: doc = fitz.open(stream=data, filetype="pdf")
    except Exception: return ""

    try:
        page_count = min(len(doc), PDF_MAX_PAGES)
        texts, total_len, digital_text_empty = [], 0, True

        for i in range(page_count):
            try: t = doc[i].get_text("text") or ""
            except Exception: t = ""
            if t.strip():
                digital_text_empty = False
                texts.append(t)
                total_len += len(t)
                if total_len > DOC_TEXT_MAX_LEN: break

        if not digital_text_empty:
            return "\n".join(texts)[:DOC_TEXT_MAX_LEN]
        
        if ocr_disabled(): return ""

        ocr_texts = []
        for i in range(page_count):
            try:
                pix = doc[i].get_pixmap(dpi=PDF_OCR_DPI)
                t = get_ocr_text(pix.tobytes("png"))
                if t:
                    ocr_texts.append(t)
                    if sum(len(x) for x in ocr_texts) > DOC_TEXT_MAX_LEN: break
            except Exception: pass
        return "\n".join(ocr_texts)[:DOC_TEXT_MAX_LEN]
    finally:
        try: doc.close()
        except Exception: pass

def _extract_docx_text(data: bytes) -> str:
    docx_lib = _get_docx()
    if docx_lib is None: return ""
    try:
        doc = docx_lib.Document(io.BytesIO(data))
        texts, total = [], 0
        for para in doc.paragraphs:
            if para.text:
                texts.append(para.text)
                total += len(para.text)
                if total > DOC_TEXT_MAX_LEN: break
        if total < DOC_TEXT_MAX_LEN:
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text:
                            texts.append(cell.text)
                            total += len(cell.text)
                            if total > DOC_TEXT_MAX_LEN: break
                    if total > DOC_TEXT_MAX_LEN: break
                if total > DOC_TEXT_MAX_LEN: break
        return "\n".join(texts)[:DOC_TEXT_MAX_LEN]
    except Exception: return ""

def _extract_xlsx_text(data: bytes) -> str:
    xl = _get_openpyxl()
    if xl is None: return ""
    try:
        wb = xl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        texts, total = [], 0
        for sheet in wb.worksheets:
            for row in sheet.iter_rows(values_only=True):
                for cell in row:
                    if cell is not None:
                        s = str(cell)
                        if s:
                            texts.append(s)
                            total += len(s)
                            if total > DOC_TEXT_MAX_LEN: break
                if total > DOC_TEXT_MAX_LEN: break
            if total > DOC_TEXT_MAX_LEN: break
        try: wb.close()
        except Exception: pass
        return "\n".join(texts)[:DOC_TEXT_MAX_LEN]
    except Exception: return ""


# ═══════════════════════════════════════════════════════════════
# 数字窗口独立扫描 —— 银行卡/身份证/手机号强力兜底
# ═══════════════════════════════════════════════════════════════
from core.patterns import (
    REGEX_PATTERNS, validate_id_card, validate_luhn,
)

def _is_valid_id_card(s: str) -> bool:
    """组合校验:GB11643 校验位 + 结构化正则(年月日 + 首位 1-9)。
    枚举路径用此校验器,防止 2-edit BFS 落到不同但合法的 GB11643 串上。"""
    if not validate_id_card(s):
        return False
    pat = REGEX_PATTERNS.get("ID_CARD")
    if pat is None:
        return True
    return pat.fullmatch(s) is not None


def _is_phone_like(s: str) -> bool:
    return (len(s) == 11 and s.isdigit()
             and s[0] == "1" and s[1] in "3456789")


def _scan_digit_windows(text: str):
    """
    提取所有"长数字窗口"(连续 13-20 位纯数字)并针对 BANK_CARD/ID_CARD/
    PHONE_NUMBER/MEDICAL_RECORD_NO 独立校验。

    动机:OCR 输出常含"6228 4893 3712 3901 69"这类空格分组,
    经 _preprocess_ocr_text 后多数可合并,但若中间混入罕见字符(如 ·/-/·)
    会破坏 _collect_findings_multipath 的常规正则。这里做"只保留数字"
    的最终兜底,极限拉升 binary_blob 层 BANK_CARD/ID_CARD 召回。

    [v8] 校验位反馈枚举:首轮直校验失败时,对 digit→digit 混淆图做有界 BFS,
         用 GB11643 / Luhn / 模板正则反验挑出"只差 1-2 位"的候选。
    """
    results = []
    seen = set()

    def _add(stype, val):
        key = (stype, val)
        if key not in seen:
            seen.add(key)
            results.append((stype, val))

    # 宽松提取所有可能的长"数字串"(允许非数字分隔符)
    # 限定字符集: 数字 + 常见分隔符(空格/-/·/中文空格)
    # 最低阈值 11 以覆盖 "138 0013 8000" 的手机号场景
    candidate_runs = re.findall(
        r"[\d\s\-·•\u3000]{11,40}",
        text,
    )

    for run in candidate_runs:
        digits_only = re.sub(r"\D", "", run)
        if not (11 <= len(digits_only) <= 20):
            continue

        # PHONE_NUMBER: 11 位,首两位 1[3-9];失败走 1-edit 枚举
        if len(digits_only) == 11:
            if _is_phone_like(digits_only):
                _add("PHONE_NUMBER", digits_only)
                continue
            fix = _enumerate_digit_fixes(digits_only, _is_phone_like,
                                          max_edits=1, max_candidates=32)
            if fix:
                _add("PHONE_NUMBER", fix)
                continue

        # ID_CARD: 18 位 GB11643 + 结构化正则双校验;失败走 1-edit 枚举
        # Why: 2-edit 虽能多救 1-2 条,但会撞出不同合法 GB11643 串(FP);
        #      1-edit 受联合校验器约束,可视为"无 FP"的安全增益。
        if len(digits_only) == 18:
            if _is_valid_id_card(digits_only):
                _add("ID_CARD", digits_only)
                continue
            fix = _enumerate_digit_fixes(digits_only, _is_valid_id_card,
                                          max_edits=1, max_candidates=128)
            if fix:
                _add("ID_CARD", fix)
                continue

        # BANK_CARD: 15-19 位, 首位 3-9, 宽松通过(OCR 场景不强制 Luhn)
        # [v8] 额外:16 位失败 Luhn → 1-edit 枚举替换(不删原条目,仅升级表达)
        if 15 <= len(digits_only) <= 19 and digits_only[0] in "356789":
            if len(digits_only) == 18 and _is_valid_id_card(digits_only):
                continue
            chosen = digits_only
            if len(digits_only) == 16 and not validate_luhn(digits_only):
                fix = _enumerate_digit_fixes(digits_only, validate_luhn,
                                              max_edits=1, max_candidates=64)
                if fix:
                    chosen = fix
            _add("BANK_CARD", chosen)

    # 也要单独做纯 11 位手机号窗口兜底(长窗口中可能有连续11位子串)
    for m in re.finditer(r"(?<!\d)1[3-9]\d{9}(?!\d)", text):
        _add("PHONE_NUMBER", m.group())

    # MEDICAL_RECORD_NO: MR + 9 位数字;扩展前缀容错 M→{M,m},R→{R,B,D,P,r,b,d,p}。
    # Why: OCR 常把 R 识成 B/D/P(字形相近);M 相对稳定。保留 M 首字母避免 FP 爆炸。
    for m in re.finditer(r"\b[Mm][RrBbDdPp]\s*(\d{9})\b", text):
        _add("MEDICAL_RECORD_NO", "MR" + m.group(1))

    return results


# ═══════════════════════════════════════════════════════════════
# 核心引擎：多路径扫描去重 (防止单点失效)
# ═══════════════════════════════════════════════════════════════
def _collect_findings_multipath(text, record_id, table_name, field_col_name,
                                 db_type, db_name):
    """
    核心战术：通过四条路径(原样、修数字、修汉字、双修)提取。
    任一路径命中且符合 patterns.py 规范的数据，都会被合并去重。
    极大缓解 OCR 极其恶劣情况下的漏报问题。

    [v5.2 加强] 追加两条独立扫描路径:
      ⑤ 数字窗口兜底 —— BANK_CARD/ID_CARD/PHONE_NUMBER 的 last-resort
      ⑥ OCR 文本内部再做一次递归解码(处理 OCR 识别出 base64/hex 明文)
    """
    if not text:
        return []

    paths = [text]

    corrected_digit = correct_ocr_text(text)
    if corrected_digit != text:
        paths.append(corrected_digit)

    corrected_name = _name_char_retry_text(text)
    if corrected_name != text:
        paths.append(corrected_name)

    if corrected_digit != text:
        corrected_both = _name_char_retry_text(corrected_digit)
        if corrected_both not in (text, corrected_digit, corrected_name):
            paths.append(corrected_both)

    # Why: OCR 常在高清扫描件里把连续数字拆成 "6228 4819 2442 3354 04"
    # (银行卡按 4 位分组印刷),或把地址数字段切成 "1353" → "135 3"。
    # 去内部 ASCII/全角空格的副本让 BANK_CARD/ID_CARD/ADDRESS 正则能连上。
    # Luhn/GB11643/USCC 校验位在下游把误拼接挡掉,所以拼错的概率很低。
    stripped = re.sub(r"[ \t\u3000]+", "", text)
    if stripped and stripped not in paths:
        paths.append(stripped)

    findings = []
    seen = set()

    def _emit(stype, val):
        key = (stype, val)
        if key in seen:
            return
        seen.add(key)
        findings.append(_make_finding(
            db_type, db_name, table_name, field_col_name,
            record_id, stype, val,
        ))

    # 主扫描路径:原文 + 修数字 + 修汉字 + 双修
    for scan_text in paths:
        for stype, val in extract_sensitive_from_value(scan_text):
            _emit(stype, val)

    # [兜底 ⑤] 数字窗口独立扫描(所有路径合并)
    for scan_text in paths:
        for stype, val in _scan_digit_windows(scan_text):
            _emit(stype, val)

    # [兜底 ⑥] OCR 文本内部可能出现的 base64/hex 编码串(罕见但存在)
    try:
        from scanners.encoded import decode_recursive as _dec
        # 对 OCR 文本中长度 > 40 的 token 逐个尝试解码
        for token in re.findall(r"[A-Za-z0-9+/=]{40,}", text):
            decoded, chain = _dec(token)
            if chain and decoded != token:
                for stype, val in extract_sensitive_from_value(decoded):
                    _emit(stype, val)
    except Exception:
        pass

    return findings

def _collect_findings_simple(text, record_id, table_name, field_col_name,
                              db_type, db_name):
    """针对无需纠错的结构化文档提取文本。"""
    if not text: return []
    findings = []
    seen = set()
    for stype, val in extract_sensitive_from_value(text):
        key = (stype, val)
        if key not in seen:
            seen.add(key)
            findings.append(_make_finding(
                db_type, db_name, table_name, field_col_name,
                record_id, stype, val,
            ))
    return findings


# ═══════════════════════════════════════════════════════════════
# 对外统一入口
# ═══════════════════════════════════════════════════════════════
def scan_blob_data(blob_bytes, record_id, table_name, field_col_name,
                   db_type, db_name):
    """
    BLOB / 文档 / 图像 字段扫描统一入口。
    供 encoded.py 和底层核心引擎调用。

    [v5.2] 识别流程新增套娃解码 + 字节级文本兜底,
    确保 binary_blob 形态下任何一种编码嵌套都不会漏扫。
    """
    data = _normalize_to_bytes(blob_bytes)
    if not data:
        return []

    # [新增] 套娃解码:BLOB 字段可能是 base64/hex 包裹的图片
    ftype = _sniff_file_type(data)
    if ftype == "unknown":
        unwrapped = _unwrap_encoded_blob(data)
        if unwrapped is not data:
            data = unwrapped
            ftype = _sniff_file_type(data)

    if ftype == "pdf":
        raw_text = _extract_pdf_text(data)
        return _collect_findings_simple(raw_text, record_id, table_name, field_col_name, db_type, db_name)

    if ftype == "docx":
        raw_text = _extract_docx_text(data)
        return _collect_findings_simple(raw_text, record_id, table_name, field_col_name, db_type, db_name)

    if ftype == "xlsx":
        raw_text = _extract_xlsx_text(data)
        return _collect_findings_simple(raw_text, record_id, table_name, field_col_name, db_type, db_name)

    # 图像或未知格式 -> 触发高规格容错处理: OCR + 预处理 + 多路径扫描
    findings = []
    seen_keys = set()

    def _merge(hits):
        for h in hits:
            key = (h["sensitive_type"], h["extracted_value"])
            if key in seen_keys:
                continue
            seen_keys.add(key)
            findings.append(h)

    # 主路径:OCR + 预处理 + 多路径正则
    if not ocr_disabled():
        raw_text = get_ocr_text(data)
        if raw_text:
            processed = _preprocess_ocr_text(raw_text)
            _merge(_collect_findings_multipath(
                processed, record_id, table_name, field_col_name,
                db_type, db_name,
            ))
            # 原始 OCR 文本(未预处理)也扫一次,防止预处理改错了格式
            if processed != raw_text:
                _merge(_collect_findings_multipath(
                    raw_text, record_id, table_name, field_col_name,
                    db_type, db_name,
                ))

    # [新增] 字节级文本兜底:BLOB 里可能混有明文字段(JSON/CSV/裸中文)
    # 无论 OCR 是否成功,都跑一次字节级 utf-8/gbk 解码扫描,命中会并入结果。
    byte_text = _bytes_as_text_fallback(data)
    if byte_text:
        _merge(_collect_findings_multipath(
            byte_text, record_id, table_name, field_col_name,
            db_type, db_name,
        ))

    return findings

# 别名: 兼容旧版本代码中的调用习惯
scan_blob_field = scan_blob_data
